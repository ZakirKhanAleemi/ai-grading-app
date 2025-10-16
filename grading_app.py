import streamlit as st
import pandas as pd
import google.generativeai as genai
import docx
import zipfile
import os
import io
import json
from PIL import Image, UnidentifiedImageError
import tempfile
import time
from google.api_core import exceptions
import fitz  # PyMuPDF
import py7zr # NEW: Library for 7z files
import rarfile # For .rar files


# --- Page Configuration (MUST BE THE FIRST STREAMLIT COMMAND) ---
st.set_page_config(layout="wide", page_title="AI Grading Assistant")

# --- Core Helper Functions ---

def extract_text_from_docx(docx_file_stream):
    """Reads text from a .docx file stream."""
    try:
        doc = docx.Document(docx_file_stream)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading .docx file: {e}"

def extract_text_from_pdf(pdf_file_stream):
    """Reads text from a .pdf file stream."""
    try:
        doc = fitz.open(stream=pdf_file_stream, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        return f"Error reading .pdf file: {e}"

def extract_images_from_docx(docx_file_stream):
    """Extracts images from a .docx file stream and returns them as PIL Image objects."""
    images = []
    try:
        doc = docx.Document(docx_file_stream)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                try:
                    image = Image.open(io.BytesIO(image_data))
                    images.append(image)
                except UnidentifiedImageError:
                    pass  # Skips unsupported image formats
    except Exception as e:
        st.error(f"Error processing a DOCX file for images: {e}")
    return images

def read_code_from_ipynb(file_path):
    """Reads and concatenates all code cells from a .ipynb file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            notebook = json.load(f)
            code_cells = [cell['source'] for cell in notebook['cells'] if cell['cell_type'] == 'code']
            return "\n".join(["".join(cell) for cell in code_cells])
    except Exception as e:
        st.warning(f"Could not read code from notebook {os.path.basename(file_path)}: {e}")
        return ""

@st.cache_resource
def get_gemini_model(_api_key):
    """Finds and configures the best available Gemini model for the user's key."""
    try:
        genai.configure(api_key=_api_key)
        model_name = None
        
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]

        # Define search priorities, from most to least capable
        priority_order = ['gemini-1.5-pro', 'gemini-pro', 'gemini-pro-vision']

        for priority in priority_order:
            for model in available_models:
                if priority in model:
                    model_name = model
                    break
            if model_name:
                break
        
        if not model_name:
            st.error("Could not find a suitable Gemini 'pro' model with your API key. Please ensure your key has access to a compatible model.")
            return None, None
        
        st.info(f"Successfully connected. Using Gemini model: `{model_name}`")
        return genai.GenerativeModel(model_name), model_name
    except Exception as e:
        st.error(f"Error configuring the Gemini API. Please check your API key. Details: {e}")
        return None, None

def grade_submission_with_gemini(model, context, rubric, code_text, report_text, images, student_folder_name, strictness_instruction):
    """Sends the student's work to the Gemini API for grading."""
    if not model:
        st.error("Grading failed: Gemini model is not available.")
        return None

    prompt_parts = [
        "You are an expert university programming tutor grading an assignment.",
        strictness_instruction,
        "Your goal is to provide a fair, detailed evaluation based *only* on the provided rubric and assignment context.",
        f"\n**Assignment Context:**\n{context}", f"\n**Grading Rubric:**\n```\n{rubric}\n```",
        f"\n**Student Submission from folder: '{student_folder_name}'**",
        "\n---", "\n**Report Text:**", report_text, "\n---",
        "\n**Python Code:**", "```python", code_text, "```", "\n---"
    ]

    if images:
        prompt_parts.append("\nThe following image(s) (e.g., flowcharts) were extracted from the report for evaluation:")
        prompt_parts.extend(images)
    else:
        prompt_parts.append("\nNo images (e.g., flowcharts) were found in the document.")

    prompt_parts.extend([
        "\n---", "\n**Your Task:**",
        "1. Evaluate the submission against each criterion in the rubric.",
        "2. For each criterion, provide a score, the maximum possible score, and a concise justification (1-2 sentences).",
        "3. Provide a single paragraph of overall constructive feedback. Keep it simple and direct (2-3 sentences).",
        "4. **You MUST respond ONLY with a single JSON object.** Do not include any text or markdown formatting before or after the JSON object.",
        """
        The JSON object must follow this exact structure:
        {
          "grading_summary": [
            { "criterion": "<Name>", "score": <num>, "max_score": <num>, "justification": "<reason>" }
          ],
          "overall_feedback": "<paragraph>"
        }
        """
    ])

    max_retries = 3
    delay = 5
    for attempt in range(max_retries):
        try:
            response = model.generate_content(prompt_parts, request_options={'timeout': 180})
            cleaned_response_text = response.text.strip().replace("```json", "").replace("```", "")
            return json.loads(cleaned_response_text)
        except exceptions.ResourceExhausted as e:
            if attempt < max_retries - 1:
                st.warning(f"Rate limit hit for '{student_folder_name}'. Retrying in {delay} seconds...")
                time.sleep(delay)
                delay *= 2
            else:
                st.error(f"Failed to grade '{student_folder_name}' due to persistent rate limits. Error: {e}")
                return None
        except Exception as e:
            st.error(f"An unexpected error occurred for '{student_folder_name}': {e}")
            return None
    return None

def find_and_grade_assignments(zip_file_object, rubric, context, model, strictness_instruction):
    """Finds and grades all student submissions from a zip file, handling nested archives."""
    all_results = []
    with tempfile.TemporaryDirectory() as temp_dir:
        status_text = st.empty()
        
        # --- Step 1: Initial Extraction ---
        try:
            with zipfile.ZipFile(zip_file_object, 'r') as zf:
                zf.extractall(temp_dir)
        except Exception as e:
            st.error(f"Failed to extract the main zip file: {e}")
            return []

        # --- Step 2: Robust Recursive Unzipping for BOTH .zip and .7z ---
        status_text.text("Unzipping nested archives...")
        while True:
            archives_found = []
            for dirpath, _, filenames in os.walk(temp_dir):
                for filename in filenames:
                    if filename.lower().endswith(('.zip', '.7z')):
                        archives_found.append(os.path.join(dirpath, filename))
            
            if not archives_found:
                break

            for archive_path in archives_found:
                extract_path = os.path.dirname(archive_path)
                try:
                    if archive_path.lower().endswith('.zip'):
                        with zipfile.ZipFile(archive_path, 'r') as zf:
                            zf.extractall(extract_path)
                    elif archive_path.lower().endswith('.7z'):
                        with py7zr.SevenZipFile(archive_path, mode='r') as z:
                            z.extractall(path=extract_path)
                    
                    os.remove(archive_path)
                except Exception as e:
                    st.warning(f"Could not extract nested archive '{os.path.basename(archive_path)}'. It might be corrupted. Error: {e}")
                    try:
                        os.remove(archive_path)
                    except OSError:
                        pass
        
        student_submission_dirs = [
            os.path.join(temp_dir, d) for d in os.listdir(temp_dir)
            if os.path.isdir(os.path.join(temp_dir, d))
        ]
        
        progress_bar = st.progress(0)
        total_assignments = len(student_submission_dirs)
        if total_assignments == 0:
            st.warning("No student folders were found in the zip file.")
            return []

        for i, student_dir_path in enumerate(student_submission_dirs):
            student_folder_name = os.path.basename(student_dir_path)
            status_text.text(f"Grading assignment {i+1}/{total_assignments}: {student_folder_name}")

            try:
                py_file, ipynb_file, docx_file, pdf_file = None, None, None, None
                for dirpath, _, filenames in os.walk(student_dir_path):
                    if not py_file: py_file = next((os.path.join(dirpath, f) for f in filenames if f.endswith('.py')), None)
                    if not ipynb_file: ipynb_file = next((os.path.join(dirpath, f) for f in filenames if f.endswith('.ipynb')), None)
                    if not docx_file: docx_file = next((os.path.join(dirpath, f) for f in filenames if f.endswith('.docx')), None)
                    if not pdf_file: pdf_file = next((os.path.join(dirpath, f) for f in filenames if f.endswith('.pdf')), None)

                report_text, code_text, images = "", "", []

                if docx_file:
                    with open(docx_file, 'rb') as f:
                        docx_stream = io.BytesIO(f.read())
                        report_text = extract_text_from_docx(docx_stream)
                        docx_stream.seek(0)
                        images = extract_images_from_docx(docx_stream)
                elif pdf_file:
                    with open(pdf_file, 'rb') as f:
                        report_text = extract_text_from_pdf(f.read())

                if py_file:
                    with open(py_file, 'r', encoding='utf-8', errors='ignore') as f:
                        code_text = f.read()
                elif ipynb_file:
                    code_text = read_code_from_ipynb(ipynb_file)
                else:
                    code_text = report_text

                if not report_text and not code_text:
                    st.warning(f"Skipping '{student_folder_name}': Could not find any content to grade.")
                    continue

                gemini_result = grade_submission_with_gemini(model, context, rubric, code_text, report_text, images, student_folder_name, strictness_instruction)

                if gemini_result:
                    total_score = sum(item.get('score', 0) for item in gemini_result.get('grading_summary', []))
                    max_score = sum(item.get('max_score', 0) for item in gemini_result.get('grading_summary', []))

                    flat_result = {
                        "Student Folder": student_folder_name, "Total Score": total_score,
                        "Max Score": max_score, "Overall Feedback": gemini_result.get('overall_feedback', 'N/A')
                    }
                    for item in gemini_result.get('grading_summary', []):
                        criterion = item.get('criterion', 'Unknown Criterion')
                        flat_result[f"{criterion} Score"] = item.get('score', 'N/A')
                        flat_result[f"{criterion} Justification"] = item.get('justification', 'N/A')
                    all_results.append(flat_result)

            except Exception as e:
                st.error(f"Failed to process assignment in folder '{student_folder_name}': {e}")
            
            progress_bar.progress((i + 1) / total_assignments)
        
        status_text.success("Grading complete!")
    return all_results

# --- Streamlit UI ---

st.title("👨‍🏫 AI-Powered Assignment Grader")
st.markdown("This tool grades assignments in a `.zip` file. It's smart enough to find code in `.py`, `.ipynb`, `.docx`, and `.pdf` files, even in nested zip or 7z archives.")

with st.sidebar:
    st.header("⚙️ Configuration")
    api_key = st.text_input("Enter your Gemini API Key:", type="password")

    # --- NEW: Grading Strictness Selector ---
    st.header("Grading Style")
    strictness_level = st.selectbox(
        "Select Grading Strictness:",
        ("Lenient", "Standard", "Strict", "Critical")
    )

    st.header("1. Define the Rubric")
    updated_rubric = """

"""
    rubric = st.text_area("Paste the grading rubric here:", value=updated_rubric, height=300)

    assignment_context = st.text_area(
        "Provide context for the assignment:",
        value=""
    )

    st.header("2. Upload Assignments")
    uploaded_file = st.file_uploader("Upload a ZIP file containing all student folders:", type="zip")

st.header("📝 Grading Results")

if 'results_df' not in st.session_state:
    st.session_state.results_df = None

if st.sidebar.button("Grade Assignments", type="primary"):
    if not api_key: st.error("Please enter your Gemini API Key.")
    elif not rubric: st.error("Please provide a grading rubric.")
    elif not uploaded_file: st.error("Please upload a zip file.")
    else:
        # --- NEW: Map selection to an instruction ---
        strictness_map = {
            "Lenient": "INSTRUCTION: Grade this submission leniently. Focus on effort and completion. Be generous with partial marks for genuine attempts, even if there are errors.",
            "Standard": "INSTRUCTION: Grade this submission fairly and accurately according to the rubric.",
            "Strict": "INSTRUCTION: Grade this submission strictly. Adherence to the rubric is critical. Deduct points for minor errors, logical flaws, or poor code quality.",
            "Critical": "INSTRUCTION: Grade this submission with a highly critical eye, as you would for a final-year project. Pay close attention to code efficiency, robustness, and professional documentation standards. Be very strict with scoring."
        }
        strictness_instruction = strictness_map[strictness_level]

        model, model_name = get_gemini_model(api_key)
        if model:
            with st.spinner(f"Grading in progress using `{model_name}` with '{strictness_level}' strictness... This may take several minutes."):
                results = find_and_grade_assignments(uploaded_file, rubric, assignment_context, model, strictness_instruction)
                if results:
                    st.session_state.results_df = pd.DataFrame(results)
                    st.success("Grading process completed successfully!")
                else:
                    st.session_state.results_df = None
                    st.warning("The grading process finished, but no results were generated.")

if st.session_state.results_df is not None:
    st.dataframe(st.session_state.results_df)

    @st.cache_data
    def convert_df_to_csv(df): return df.to_csv(index=False).encode('utf-8')
    @st.cache_data
    def convert_df_to_excel(df):
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Grades')
        return output.getvalue()

    csv_data = convert_df_to_csv(st.session_state.results_df)
    excel_data = convert_df_to_excel(st.session_state.results_df)

    col1, col2 = st.columns(2)
    with col1:
        st.download_button("📥 Download Results as CSV", csv_data, 'grading_report.csv', 'text/csv')
    with col2:
        st.download_button("📥 Download Results as Excel", excel_data, 'grading_report.xlsx')
else:
    st.info("Results will be displayed here after the grading process is complete.")

